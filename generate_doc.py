import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def main():
    doc = Document()

    # --- Title Page ---
    title = doc.add_heading('Analisis Sistem Komprehensif\nALS Energy Monitoring', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Dokumen Analisis Arsitektur, Basis Data, dan Logika Sistem')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # --- 1. Arsitektur Sistem ---
    add_heading(doc, '1. Arsitektur Sistem (System Architecture)', 1)
    
    p = doc.add_paragraph('Sistem ALS Energy dibangun dengan pendekatan arsitektur ')
    p.add_run('Modular (Monolithic-Modular)').bold = True
    p.add_run(', memisahkan komponen Frontend dan Backend namun disajikan secara terintegrasi melalui ')
    p.add_run('Flask Blueprints.').bold = True
    
    doc.add_paragraph('Teknologi Utama:', style='List Bullet')
    ul1 = doc.add_paragraph(style='List Bullet 2')
    ul1.add_run('Backend:').bold = True
    ul1.add_run(' Python (Flask Framework).')
    ul2 = doc.add_paragraph(style='List Bullet 2')
    ul2.add_run('Database:').bold = True
    ul2.add_run(' MySQL (dikoneksikan via mysql-connector-python).')
    ul3 = doc.add_paragraph(style='List Bullet 2')
    ul3.add_run('Frontend:').bold = True
    ul3.add_run(' HTML5 murni, Vanilla JavaScript, CSS, TailwindCSS (CDN), dan Chart.js untuk visualisasi data.')
    
    doc.add_paragraph('Pengamanan Endpoint API (JWT):')
    doc.add_paragraph('Sistem tidak menggunakan sesi statis (cookies), melainkan menggunakan sistem otentikasi JSON Web Tokens (JWT). Setiap akses menuju API tertutup (Private API) seperti mengambil riwayat atau mengubah pengaturan wajib melampirkan Token JWT di header yang validasinya dilakukan oleh Middleware terpusat (@token_required).')

    # --- 2. Struktur Menu & Antarmuka ---
    add_heading(doc, '2. Struktur Menu & Antarmuka (Menu Structure)', 1)
    
    doc.add_paragraph('Sistem diakses melalui panel admin tunggal yang terdiri dari menu-menu berikut:')
    menus = [
        ("Dashboard (index.html)", "Menampilkan ringkasan sistem, persentase baterai (dummy), total sensor aktif, ringkasan notifikasi terbaru, dan log singkat."),
        ("Monitoring (monitoring.html)", "Menyajikan data Real-Time menggunakan Gauge Charts (Tegangan, Arus, Daya Aktif) dan Line Chart dinamis yang memperbarui data setiap 3 detik."),
        ("Riwayat Data (riwayat_data.html)", "Menyajikan seluruh data historis dari database dalam format tabel yang dilengkapi dengan fitur Paginasi Server-side (Navigasi Halaman) dan Filter Batas Tampil (10/20/50/100 data)."),
        ("Prediksi AI (prediksi.html)", "Halaman khusus analisis prediktif untuk memproyeksikan sisa hari penggunaan listrik berdasarkan nominal uang yang diinputkan."),
        ("Log Sistem & Notifikasi", "Mencatat jejak audit aktivitas yang terjadi di backend (login, setelan diubah) dan riwayat pengiriman peringatan Telegram."),
        ("Pengaturan (pengaturan.html)", "Antarmuka untuk mengatur konfigurasi Bot Telegram (Token, Chat ID), Kuota Energi (Ambang Batas), dan Tarif Dasar Listrik (Rp/kWh).")
    ]
    for menu, desc in menus:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(menu + ': ').bold = True
        p.add_run(desc)

    # --- 3. Struktur Database ---
    add_heading(doc, '3. Struktur Basis Data (Database Schema)', 1)
    doc.add_paragraph('Database yang digunakan bernama `als_energy` yang terdiri dari 5 tabel fungsional utama:')
    
    tables = [
        ("admins", "Tabel untuk menyimpan kredensial login. Kolom utama: id, username, password_hash."),
        ("sensor_data", "Tabel transaksi utama (Big Data). Menyimpan seluruh metrik kelistrikan. Kolom utama: timestamp, mesin_id, volt, arus, daya, energi, frekuensi, pf."),
        ("settings", "Tabel konfigurasi fleksibel (Key-Value Pair). Kolom utama: setting_key, setting_value. Digunakan untuk menyimpan parameter bot dan tarif."),
        ("system_events", "Tabel untuk pencatatan log (Audit Trail). Kolom utama: timestamp, event_type, description."),
        ("notifications", "Tabel pencatatan riwayat pengiriman pesan peringatan (Alerts). Kolom utama: timestamp, message, status.")
    ]
    for t_name, t_desc in tables:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(t_name).bold = True
        p.add_run(' - ' + t_desc)

    # --- 4. Use Case & Aktor Sistem ---
    doc.add_page_break()
    add_heading(doc, '4. Model Use Case', 1)
    doc.add_paragraph('Sistem ini didesain untuk satu entitas aktor utama, yaitu Administrator. Interaksi use case mencakup:')
    
    use_cases = [
        "Login dan Logout ke dalam sistem (Otentikasi).",
        "Melihat status kelistrikan dua mesin yang berbeda secara Real-Time.",
        "Mengekspor dan mencetak riwayat data penggunaan listrik.",
        "Mengeksekusi simulasi Prediksi AI (Melihat ramalan konsumsi energi berdasar anggaran/budget).",
        "Menghidupkan/Mematikan integrasi Bot Telegram.",
        "Memasukkan batas kritis (Alert Trigger) yang akan memicu Telegram me-notifikasi ponsel Admin secara otomatis."
    ]
    for uc in use_cases:
        doc.add_paragraph(uc, style='List Number')

    # --- 5. Flowchart Logika & AI ---
    add_heading(doc, '5. Penjabaran Flowchart Proses (Algoritma)', 1)
    
    add_heading(doc, 'A. Flowchart Aliran Data Sensor (Ingestion)', 2)
    doc.add_paragraph('1. Sensor ESP32 (PZEM) mendeteksi V, A, P, dan E.')
    doc.add_paragraph('2. ESP32 menembak HTTP POST Request berisi data JSON ke Endpoint Backend (/api/sensor/data).')
    doc.add_paragraph('3. Backend menerima, melakukan validasi, lalu menyimpannya ke tabel `sensor_data`.')
    doc.add_paragraph('4. Backend melakukan asinkronisasi cek ambang batas: Apakah (Kuota Energi Maksimal - Energi Terkini) <= Batas Kritis?')
    doc.add_paragraph('5. Jika Ya (Kritis): Sistem memanggil API Telegram HTTP dan menembakkan pesan darurat ke HP Admin.')
    doc.add_paragraph('6. Jika Tidak: Proses selesai tanpa peringatan.')
    
    add_heading(doc, 'B. Algoritma Prediksi (Double Exponential Smoothing)', 2)
    p = doc.add_paragraph('Sistem prediksi tidak menggunakan Machine Learning rumit yang lamban, melainkan model peramalan deret waktu (Time Series Forecasting) tingkat lanjut bernama ')
    p.add_run('Holt\'s Double Exponential Smoothing').bold = True
    p.add_run('.')
    
    doc.add_paragraph('Tahapan peramalan matematis:')
    p = doc.add_paragraph('1. Konversi Uang: Sistem menerima nominal budget, lalu membaginya dengan Tarif Dasar (Rp/kWh) untuk mendapatkan Target Total kWh.')
    p = doc.add_paragraph('2. Penarikan Riwayat: Sistem menarik seluruh daya rata-rata per hari dari database.')
    p = doc.add_paragraph('3. Inisialisasi: Parameter Alpha (0.3) untuk pelemahan level, dan Beta (0.2) untuk perhitungan kemiringan Tren (Trend).')
    p = doc.add_paragraph('4. Iterasi (Training): Model mensimulasikan perhitungan eksponensial di tiap harinya (S[t] dan B[t]) untuk mempelajari kecenderungan data (Apakah listrik makin boros atau stabil?).')
    p = doc.add_paragraph('5. Proyeksi (Forecasting): Model kemudian memproyeksikan angka pemakaian listrik untuk hari-hari ke depan. Hari-hari tersebut dihitung hingga akumulasi energi menabrak batas Target Total kWh.')
    p = doc.add_paragraph('6. Output akhir berupa angka "Estimasi Hari" dan "Jam", serta grafik prakiraan (Forecast Chart) untuk 7 hari mendatang.')

    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'Analisis_Sistem_ALS_Energy.docx')
    doc.save(output_path)
    print(f"File DOCX berhasil dibuat di: {output_path}")

if __name__ == "__main__":
    main()
